import docx

class FileChecker:
    def __init__(self, doc_path):
        self.doc = docx.Document(doc_path)

    def check_title_presence(self, expected_titles):
        """Check if specific titles exist in the document with the correct content."""
        titles_found = {title: False for title in expected_titles}
        for para in self.doc.paragraphs:
            content = para.text.strip()
            if content in titles_found:
                titles_found[content] = True
                print(f"{content} section found.")  # 打印找到的章节

                # 如果找到所有标题，提前返回
                if all(titles_found.values()):
                    return True
        return False

    def check_statement_page(self):
        """Check if the STATEMENT section exists in the document."""
        for para in self.doc.paragraphs:
            if para.text.strip() == "STATEMENT":
                print("STATEMENT section found.")
                return True

        print("STATEMENT section is missing.")
        return False

    def check_document_structure(self):
        """Main function to check the structure based on specific titles."""
        structure_checks = {
            "STATEMENT": self.check_statement_page(),  # 检查 STATEMENT 是否存在
            "ABOUT THE DOCUMENT and TARGET USERS": self.check_title_presence(["ABOUT THE DOCUMENT", "TARGET USERS"]),
            "SYMBOL DESCRIPTION and EXPLANATION OF TERMS": self.check_title_presence(["SYMBOL DESCRIPTION", "EXPLANATION OF TERMS"])
        }
        
        missing_sections = []
        for section, result in structure_checks.items():
            if not result:
                missing_sections.append(section)
                print(f"Detailed Issue: '{section}' page is missing or incorrect.")

        # 总结检查结果
        if not missing_sections:
            print("Document structure is correct.")
        else:
            print("Document structure is incomplete or incorrect.")
            print("Missing or incorrect sections:")
            for section in missing_sections:
                print(f"- {section}")

    def check_font_consistency(self, allowed_fonts=("Montserrat", "思源黑体")):
        """Check all text in the document uses one of the allowed fonts."""
        font_issues = []
        for i, para in enumerate(self.doc.paragraphs):
            if not para.text.strip():  # Skip empty paragraphs
                continue

            for run in para.runs:
                if run.text.strip() and (run.font.name is not None) and (run.font.name not in allowed_fonts):
                    font_issues.append(
                        f"Issue at paragraph {i+1}, text: '{para.text.strip()}': found font '{run.font.name}' instead of one of {allowed_fonts}"
                    )

        if font_issues:
            print("Font consistency issues found:")
            for issue in font_issues:
                print(issue)
        else:
            print("Font consistency is correct.")
