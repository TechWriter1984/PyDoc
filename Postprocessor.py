import docx
import jieba

class Postprocessor:
    """
    Handles the postprocessing phase for Word documents, such as adding supplementary sections.
    """

    def __init__(self, glossary):
        """
        Initialize the Postprocessor with a glossary of terms.
        :param glossary: List of dictionaries containing terms, expanded forms, and descriptions.
        """
        self.glossary = glossary

    def insert_explanation_of_terms(self, doc, terms):
        """
        Inserts an "EXPLANATION OF TERMS" section at the beginning of the document.
        :param doc: docx.Document object.
        :param terms: List of matched terms to insert.
        """
        # Add a new page break for the EXPLANATION OF TERMS section
        doc.add_paragraph("\n")
        explanation_paragraph = doc.add_paragraph("EXPLANATION OF TERMS")
        explanation_paragraph.style = "Heading 1"

        # Create a table to store terms and their explanations
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        # Add headers to the table
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Acronym"
        hdr_cells[1].text = "Expanded Form"
        hdr_cells[2].text = "Description"

        # Populate the table with terms
        for term in terms:
            row_cells = table.add_row().cells
            row_cells[0].text = term["acronym"]
            row_cells[1].text = term["expanded_form"]
            row_cells[2].text = term["description"]

        print("EXPLANATION OF TERMS section has been added.")

    def detect_terms(self, doc):
        """
        Detects terms from the document that match the glossary.
        :param doc: docx.Document object.
        :return: List of matched terms.
        """
        # Combine all document text
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        full_text = " ".join(full_text)

        # Perform word segmentation using jieba
        words = jieba.lcut(full_text)

        # Match detected words with glossary terms
        matched_terms = [term for term in self.glossary if term["acronym"] in words]
        return matched_terms

    def process_word_file(self, input_path, output_path):
        """
        Postprocesses the Word file to add an EXPLANATION OF TERMS section if necessary.
        :param input_path: Path to the input Word file.
        :param output_path: Path to save the processed Word file.
        """
        try:
            # Load the document
            doc = docx.Document(input_path)

            # Detect terms in the document
            matched_terms = self.detect_terms(doc)

            # If matched terms exist, add the EXPLANATION OF TERMS section
            if matched_terms:
                self.insert_explanation_of_terms(doc, matched_terms)
            else:
                print("No matching terms found. No EXPLANATION OF TERMS section added.")

            # Save the updated document
            doc.save(output_path)
            print(f"Postprocessing completed. Processed file saved at {output_path}")

        except Exception as e:
            print(f"An error occurred during postprocessing: {e}")
