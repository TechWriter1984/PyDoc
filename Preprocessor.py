import docx
import os
import shutil

class Preprocessor:

    def delete_before_heading1(self, doc):
        """
        删除文档中第一个"Heading 1"之前的所有内容。
        
        参数:
            doc (Document): 要处理的docx文档对象
        
        功能:
            1. 获取文档正文中的所有元素
            2. 遍历查找第一个"Heading 1"样式的段落
            3. 如果找到，删除该段落之前的所有内容
            4. 如果未找到，不执行任何操作
        """
        body = doc.element.body
        elements = list(body)  # 获取文档中的所有元素(段落和表格)
    
        for idx, element in enumerate(elements):
            if element.tag.endswith('p'):  # 检查是否是段落元素
                para = docx.text.paragraph.Paragraph(element, doc)
                if para.style and para.style.name == "Heading 1":
                    # 删除第一个"Heading 1"之前的所有元素
                    for e in elements[:idx]:
                        body.remove(e)
                    print("所有正文前内容已被删除。")
                    return
        print("文档中未找到Heading 1，未删除任何内容。")

    def delete_headers_footers(self, doc):
        """
        删除文档中的所有页眉和页脚内容。
        
        参数:
            doc (Document): 要处理的docx文档对象
        
        功能:
            1. 遍历文档的所有节(section)
            2. 取消页眉页脚与前一节的链接
            3. 删除每个节中页眉和页脚的所有段落
        """
        for section in doc.sections:
            # 取消页眉页脚与前一节的链接
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
    
            # 删除页眉中的所有段落
            for paragraph in section.header.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
            # 删除页脚中的所有段落
            for paragraph in section.footer.paragraphs:
                p = paragraph._element
                p.getparent().remove(p)
    
        print("所有页眉和页脚已被删除。")

    def remove_watermark(self, doc):
        """
        Try to remove watermark from the document (based on XML structure).
        Notice: This method may fail if the document structure changes.
        """
        try:
            # Access the XML content of the document and delete the watermark node
            for element in doc.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}header'):
                for watermark in element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                    element.remove(watermark)
            print("Watermark has been removed, if any.")
        except Exception as e:
            print(f"An error occurred while removing watermark: {e}")

    def process_word_file(self, input_path, output_path):
        """
        Process a Word file: delete all content before the first "Heading 1", delete headers, footers, and watermark.
        """
        # Create a backup path for the output file
        backup_path = f"{output_path}.backup.docx"
        try:
            shutil.copyfile(input_path, backup_path)
            print(f"A backup of the original file has been created at {backup_path}")

        except Exception as e:
            print(f"An error occurred while creating the backup: {e}")
            return

        try:
            doc = docx.Document(input_path)
            self.delete_before_heading1(doc)
            self.delete_headers_footers(doc)
            self.remove_watermark(doc)
            # Save the processed document to the output path
            doc.save(output_path)
            print(f"Doc preprocessing completed. The processed document is saved at {output_path}")

        except Exception as e:
            print(f"An error occurred during document preprocessing: {e}")