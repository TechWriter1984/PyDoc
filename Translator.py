import docx
import json
from urllib.parse import quote
import datetime
import hmac
import hashlib
import requests
import os
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class Translator:
    """
    Initializes the Translator class.

    This method creates an empty dictionary translated_cache that is used to cache translated text to avoid repeated translation.
    """
    def __init__(self):
        self.translated_cache = {}

    def hmac_sha256(self, key: bytes, content: str):
        """
        Calculates the HMAC-SHA256 hash of the given content using the provided key.

        Args:
            key (bytes): The secret key used for HMAC-SHA256 calculation.
            content (str): The content to be hashed.

        Returns:
            bytes: The HMAC-SHA256 hash of the content.
        """
        return hmac.new(key, content.encode("utf-8"), hashlib.sha256).digest()

    def hash_sha256(self, content: str):
        """
        Calculates the SHA256 hash of the given content.

        Args:
            content (str): The content to be hashed.

        Returns:
            str: The hexadecimal representation of the SHA256 hash.
        """
        return hashlib.sha256(content.encode("utf-8")).hexdigest()

    def norm_query(self, params):
        """
        Normalizes the query parameters into a URL-encoded string.

        Args:
            params (dict): A dictionary containing the query parameters.

        Returns:
            str: The normalized query string.
        """
        query = ""
        for key in sorted(params.keys()):
            if type(params[key]) == list:
                for k in params[key]:
                    query = query + quote(key, safe="-_.~") + "=" + quote(k, safe="-_.~") + "&"
            else:
                query = query + quote(key, safe="-_.~") + "=" + quote(params[key], safe="-_.~") + "&"
        return query[:-1].replace("+", "%20")

    def request(self, method, date, query, header, ak, sk, action, body):
        """
        Sends a request to the specified API endpoint with the given parameters.

        Args:
            method (str): The HTTP method to use for the request (e.g., "GET", "POST").
            date (datetime.datetime): The date and time of the request.
            query (dict): The query parameters for the request.
            header (dict): The HTTP headers for the request.
            ak (str): The access key ID for authentication.
            sk (str): The secret access key for authentication.
            action (str): The API action to perform.
            body (str): The body of the request (if any).

        Returns:
            dict: The JSON response from the API, or None if an error occurred.
        """
        credential = {
            "access_key_id": ak,
            "secret_access_key": sk,
            "service": "translate",
            "region": "cn-north-1",
        }
        request_param = {
            "body": body,
            "host": "translate.volcengineapi.com",
            "path": "/",
            "method": method,
            "content_type": "application/json",
            "date": date,
            "query": {"Action": action, "Version": "2020-06-01", **query},
        }
        if body is None:
            request_param["body"] = ""
        x_date = request_param["date"].strftime("%Y%m%dT%H%M%SZ")
        short_x_date = x_date[:8]
        x_content_sha256 = self.hash_sha256(request_param["body"])
        sign_result = {
            "Host": request_param["host"],
            "X-Content-Sha256": x_content_sha256,
            "X-Date": x_date,
            "Content-Type": request_param["content_type"],
        }
        signed_headers_str = ";".join(["content-type", "host", "x-content-sha256", "x-date"])
        canonical_request_str = "\n".join(
            [request_param["method"].upper(),
             request_param["path"],
             self.norm_query(request_param["query"]),
             "\n".join(
                 [
                     "content-type:" + request_param["content_type"],
                     "host:" + request_param["host"],
                     "x-content-sha256:" + x_content_sha256,
                     "x-date:" + x_date,
                 ]
             ),
             "",
             signed_headers_str,
             x_content_sha256,
             ]
        )
        hashed_canonical_request = self.hash_sha256(canonical_request_str)
        credential_scope = "/".join([short_x_date, credential["region"], credential["service"], "request"])
        string_to_sign = "\n".join(["HMAC-SHA256", x_date, credential_scope, hashed_canonical_request])
        k_date = self.hmac_sha256(credential["secret_access_key"].encode("utf-8"), short_x_date)
        k_region = self.hmac_sha256(k_date, credential["region"])
        k_service = self.hmac_sha256(k_region, credential["service"])
        k_signing = self.hmac_sha256(k_service, "request")
        signature = self.hmac_sha256(k_signing, string_to_sign).hex()
        sign_result["Authorization"] = "HMAC-SHA256 Credential={}, SignedHeaders={}, Signature={}".format(
            credential["access_key_id"] + "/" + credential_scope,
            signed_headers_str,
            signature,
        )
        header = {**header, **sign_result}

        try:
            print(f"Making request with method: {method}, url: https://{request_param['host']}{request_param['path']}, params: {request_param['query']}, headers: {header}, data: {request_param['body']}")
            r = requests.request(
                method=method,
                url="https://{}{}".format(request_param["host"], request_param["path"]),
                headers=header,
                params=request_param["query"],
                data=request_param["body"],
                proxies={},
            )
            print(f"Received response status code: {r.status_code}, content: {r.content}")
            r.raise_for_status()
            return r.json()
        except requests.exceptions.RequestException as e:
            print(f"An error occurred while making the request: {e}")
            return None

    def translate_text(self, text):
        """
        Translates the given text to English.

        Args:
            text (str): The text to be translated.

        Returns:
            str: The translated text, or an error message if the translation fails.
        """
        if text in self.translated_cache:
            return self.translated_cache[text]
        body = {
            'TargetLanguage': 'en',
            'TextList': [text],
        }
        url = "https://translate.volcengineapi.com/"
        now = datetime.datetime.now(datetime.timezone.utc)
        headers = {}
        result = self.request("POST", now, {}, headers, os.getenv('VOLC_ACCESS_KEY'), os.getenv('VOLC_SECRET_KEY'), "TranslateText", json.dumps(body))
        if result is None:
            print("Translation result is None. Check the request.")
            return "Failed to obtain translation result."
        elif 'TranslationList' in result:
            self.translated_cache[text] = result['TranslationList'][0]['Translation']
            return self.translated_cache[text]
        else:
            print(f"Unexpected result: {result}")
            return "Failed to obtain translation result."

    def insert_paragraph_after(self, para, text=None, style=None):
        """
        Inserts a new paragraph after the given paragraph with the specified text and style.

        Args:
            para (docx.text.paragraph.Paragraph): The paragraph after which the new paragraph will be inserted.
            text (str, optional): The text content of the new paragraph. Defaults to None.
            style (docx.styles.style.Style, optional): The style of the new paragraph. Defaults to None.

        Returns:
            docx.text.paragraph.Paragraph: The newly inserted paragraph object.
        """
        new_para = OxmlElement("w:p")
        para._element.addnext(new_para)
        new_paragraph = docx.text.paragraph.Paragraph(new_para, para._parent)
        if text:
            run = new_paragraph.add_run(text)
            if style and style.type == docx.enum.style.WD_STYLE_TYPE.CHARACTER:
                run.style = style
            elif style and style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH:
                new_paragraph.style = style
        return new_paragraph

    def translate_word_file(self, input_path, output_path):
            """Translate the document and add translation text below the original text while retaining styles."""
            doc = docx.Document(input_path)

            # Translate all paragraphs
            for para in doc.paragraphs:
                original_text = para.text
                if original_text:
                    translated_text = self.translate_text(original_text)
                    new_para = self.insert_paragraph_after(para, text=translated_text, style=para.style)

            # Translate all tables, using a cache to avoid duplicate translations
            translated_table_cells = {}

            # Iterate through all tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        original_text = cell.text.strip()
                        if original_text:
                            if (row, cell) not in translated_table_cells:
                                lines = original_text.splitlines()
                                translated_lines = [self.translate_text(line) for line in lines]
                                translated_table_cells[(row, cell)] = translated_lines
                            else:
                                translated_lines = translated_table_cells[(row, cell)]
                            
                            cell.text = ""
                            for line, translated_line in zip(lines, translated_lines):
                                para = cell.add_paragraph(line)
                                translated_para = cell.add_paragraph(translated_line)
                                translated_para.style = para.style

            doc.save(output_path)
            print(f"Translation completed. Document saved at: {output_path}")
