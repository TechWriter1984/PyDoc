import os
import docx
from typing import Dict, List, Optional, Any


class DocumentTester:
    """
    文档测试器类，用于检查文档的必要部件完整性。
    
    该类提供了全面检查Word文档中必要部件的功能，支持检查正文、页眉和页脚中的部件标识。
    通过验证文档中的关键标识，确保文档结构符合标准要求。
    """
    
    # 定义需要检查的文档部件
    # 键: 部件中文名称 (用于结果报告显示)
    # 值: 部件标识符 (在文档中实际搜索的文本)
    REQUIRED_PARTS = {
        "封面": "[Cover]",      # 隐藏文字标识
        "声明": "STATEMENT",    # 大写英文标识
        "关于文档": "ABOUT THIS DOCUMENT",  # 大写英文标识
        "目标读者": "TARGET USERS",        # 大写英文标识
        "符号说明": "SYMBOL DESCRIPTION",   # 大写英文标识
        "术语解释": "EXPLANATION OF TERMS",  # 大写英文标识
        "目录": "Table of Contents",       # 标准目录标题
        "正文": "[Body]"        # 隐藏文字标识
    }
    
    def __init__(self):
        """
        初始化文档测试器实例。
        
        初始化检查结果的相关属性，包括：
        - results: 存储各部件的检查状态
        - missing_parts: 记录缺失的部件
        - found_parts: 记录找到的部件
        """
        self.results = {}
        self.missing_parts = []
        self.found_parts = []
    
    def check_document_parts(self, file_path: str) -> Dict[str, Any]:
        """
        检查文档的必要部件是否完整。
        
        该方法会验证文档是否包含所有必要的部件，如封面、声明、目录等。
        支持检查文档正文、页眉和页脚中的部件标识，包括隐藏文字类型的标识。
        
        Args:
            file_path: 文档的绝对路径
            
        Returns:
            包含检查结果的字典，格式为 {
                "status": "success" 或 "error",
                "message": 结果消息,
                "missing_parts": 缺失的部件列表,
                "found_parts": 找到的部件列表,
                "total_required": 总需要检查的部件数,
                "total_found": 实际找到的部件数
            }
        
        Raises:
            ValueError: 当提供的路径无效或不是绝对路径时
            FileNotFoundError: 当文档文件不存在时
            Exception: 当处理文档时发生其他错误
        """
        # 验证文件路径
        if not file_path:
            raise ValueError("文件路径不能为空")
        
        if not os.path.isabs(file_path):
            raise ValueError(f"提供的路径不是绝对路径: {file_path}")
        
        try:
            # 检查文件是否存在
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"错误：文件不存在 - {file_path}")
            
            # 检查文件类型
            if not file_path.lower().endswith('.docx'):
                raise ValueError(f"仅支持.docx格式的文档: {file_path}")
            
            # 重置结果
            self.results = {part: False for part in self.REQUIRED_PARTS.keys()}
            self.missing_parts = []
            self.found_parts = []
            
            print(f"正在检查文档: {os.path.basename(file_path)}")
            
            # 加载文档
            doc = docx.Document(file_path)
            
            print("检查正文段落...")
            # 检查所有段落
            for para in doc.paragraphs:
                self._check_paragraph(para)
            
            print("检查页眉和页脚...")
            # 检查所有页眉和页脚（可能包含隐藏文字）
            for section in doc.sections:
                # 检查页眉
                for header_para in section.header.paragraphs:
                    self._check_paragraph(header_para, is_hidden=True)
                
                # 检查页脚
                for footer_para in section.footer.paragraphs:
                    self._check_paragraph(footer_para, is_hidden=True)
            
            # 生成检查结果
            for part, is_found in self.results.items():
                if is_found:
                    self.found_parts.append(part)
                else:
                    self.missing_parts.append(part)
            
            # 生成报告消息
            if not self.missing_parts:
                message = "检查完成：所有必要部件都已找到。"
            else:
                message = f"检查完成：发现{len(self.missing_parts)}个缺失部件。请添加这些部件。"
            
            return {
                "status": "success",
                "message": message,
                "missing_parts": self.missing_parts,
                "found_parts": self.found_parts,
                "total_required": len(self.REQUIRED_PARTS),
                "total_found": len(self.found_parts),
                "file_path": file_path  # 添加文件路径到结果中
            }
            
        except (ValueError, FileNotFoundError) as e:
            # 处理特定异常
            return {
                "status": "error",
                "message": str(e),
                "missing_parts": [],
                "found_parts": [],
                "total_required": len(self.REQUIRED_PARTS),
                "total_found": 0,
                "file_path": file_path
            }
        except docx.opc.exceptions.PackageNotFoundError:
            # 处理文档格式错误
            return {
                "status": "error",
                "message": f"无效的Word文档格式: {file_path}",
                "missing_parts": [],
                "found_parts": [],
                "total_required": len(self.REQUIRED_PARTS),
                "total_found": 0,
                "file_path": file_path
            }
        except PermissionError:
            # 处理文件权限错误
            return {
                "status": "error",
                "message": f"无权限访问文件: {file_path}",
                "missing_parts": [],
                "found_parts": [],
                "total_required": len(self.REQUIRED_PARTS),
                "total_found": 0,
                "file_path": file_path
            }
        except Exception as e:
            # 处理其他异常
            return {
                "status": "error",
                "message": f"检查过程中发生错误：{str(e)}",
                "missing_parts": [],
                "found_parts": [],
                "total_required": len(self.REQUIRED_PARTS),
                "total_found": 0,
                "file_path": file_path
            }
    
    def _check_paragraph(self, para: docx.text.paragraph.Paragraph, is_hidden: bool = False) -> None:
        """
        检查单个段落是否包含某个文档部件标识。
        
        Args:
            para: 要检查的段落
            is_hidden: 是否为隐藏文字部分
        """
        text = para.text.strip()
        
        # 检查每个必要部件
        for part_name, part_identifier in self.REQUIRED_PARTS.items():
            # 如果已经找到该部件，跳过检查
            if self.results[part_name]:
                continue
            
            # 检查是否匹配
            if part_identifier.startswith("[") and part_identifier.endswith("]"):
                # 隐藏文字类型 ([Cover], [Body])
                # 对于隐藏文字，我们需要精确匹配
                if text == part_identifier:
                    # 隐藏文字可能在任何位置，包括页眉页脚
                    self.results[part_name] = True
            else:
                # 大写英文标题类型 (STATEMENT, ABOUT THIS DOCUMENT 等)
                # 对于大写英文标题，我们需要精确匹配
                if text == part_identifier:
                    self.results[part_name] = True
                    
        # 特殊检查：对于目录部分，还需要检查可能的变体
        if not self.results["目录"]:
            # 检查可能的目录标题变体
            directory_variants = ["Table of Contents", "TABLE OF CONTENTS", "Contents", "CONTENTS"]
            if text in directory_variants:
                self.results["目录"] = True
    
    def generate_report(self, result: Dict[str, any]) -> str:
        """
        生成详细的检查报告。
        
        Args:
            result: 检查结果字典
            
        Returns:
            格式化的报告字符串
        """
        report = []
        report.append("=" * 80)
        report.append("文档部件完整性检查报告".center(76))
        report.append("=" * 80)
        report.append(f"状态: {'✓ 成功' if result['status'] == 'success' else '✗ 失败'}")
        report.append(f"消息: {result['message']}")
        report.append(f"检查统计: {result['total_found']}/{result['total_required']} 部件已找到")
        
        # 计算完成百分比
        if result['total_required'] > 0:
            completion_rate = (result['total_found'] / result['total_required']) * 100
            report.append(f"完成率: {completion_rate:.1f}%")
        
        report.append("-" * 80)
        
        if result['found_parts']:
            report.append("找到的部件:")
            for part in result['found_parts']:
                identifier = self.REQUIRED_PARTS[part]
                report.append(f"  ✓ {part.ljust(12)} - {identifier}")
            report.append("-" * 80)
        
        if result['missing_parts']:
            report.append("缺失的部件:")
            for part in result['missing_parts']:
                identifier = self.REQUIRED_PARTS[part]
                report.append(f"  ✗ {part.ljust(12)} - {identifier}")
            report.append("-" * 80)
            report.append("请根据上述信息添加缺失的部件以确保文档完整性。")
        else:
            report.append("恭喜！文档包含所有必要的部件。")
        
        report.append("-" * 80)
        report.append("说明:")
        report.append("  1. 此检查仅验证文档部件的存在性")
        report.append("  2. 尚未根据文档类型(产品说明书、安装手册等)测试内容和目录设计")
        report.append("  3. 后续将根据文档类型进一步扩展检查功能")
        report.append("=" * 80)
        
        return "\n".join(report)

# 为了方便直接使用，提供一个函数接口
def check_document_parts(file_path: str) -> Dict[str, Any]:
    """
    检查文档的必要部件是否完整的便捷函数。
    
    该函数是DocumentTester类的简化接口，提供了检查文档部件完整性的快捷方法。
    它会创建DocumentTester实例，执行检查，并打印详细报告。
    
    Args:
        file_path: 文档的绝对路径
        
    Returns:
        包含检查结果的字典，包括状态、消息、找到的部件和缺失的部件
        
    Raises:
        ValueError: 当提供的路径无效或不是绝对路径时
        FileNotFoundError: 当文档文件不存在时
    """
    tester = DocumentTester()
    result = tester.check_document_parts(file_path)
    report = tester.generate_report(result)
    print(report)
    return result

if __name__ == "__main__":
    # 示例用法
    import sys
    if len(sys.argv) > 1:
        check_document_parts(sys.argv[1])
    else:
        print("用法: python doc_tester.py <文档路径>")
